# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # 添加標題
    title = doc.add_heading('特徵生成器(03_feature_generator)說明文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 功能概述
    doc.add_heading('1. 功能概述', 1)
    
    # 1.1 核心功能
    doc.add_heading('1.1 核心功能', 2)
    core_features = [
        '量能特徵生成',
        '波動特徵計算',
        '趨勢特徵分析',
        '技術指標整合',
        '特徵品質監控',
        '自動報告生成'
    ]
    for feature in core_features:
        p = doc.add_paragraph()
        p.add_run('• ' + feature)
    
    # 1.2 特徵類別
    doc.add_heading('1.2 特徵類別', 2)
    
    # 量能類特徵
    doc.add_heading('1. 量能類特徵', 3)
    volume_features = [
        '量比：當日成交量/N日平均成交量',
        '量增率：當日成交量相對前一日變化率',
        '量能趨勢：短期均量/長期均量'
    ]
    for feature in volume_features:
        p = doc.add_paragraph()
        p.add_run('- ' + feature)
    
    # 波動類特徵
    doc.add_heading('2. 波動類特徵', 3)
    volatility_features = [
        '日內波動率：(最高價-最低價)/開盤價',
        '振幅：(最高價-最低價)/前一日收盤價',
        '漲跌幅：價格變動百分比',
        '波動率趨勢：短期波動率/長期波動率'
    ]
    for feature in volatility_features:
        p = doc.add_paragraph()
        p.add_run('- ' + feature)
    
    # 趨勢類特徵
    doc.add_heading('3. 趨勢類特徵', 3)
    trend_features = [
        '趨勢強度：收盤價相對移動平均偏離程度',
        '通道寬度變化：布林通道寬度變化率',
        '趨勢動能：價格變動速度',
        '趨勢持續性：趨勢方向的持續性指標'
    ]
    for feature in trend_features:
        p = doc.add_paragraph()
        p.add_run('- ' + feature)
    
    # 技術類特徵
    doc.add_heading('4. 技術類特徵', 3)
    technical_features = [
        'KD_差值: K線與D線之間的差距值',
        'RSI_動能: RSI值的變化速度',
        'MACD_動能: MACD柱狀圖的變化率',
        '均線糾結度: 移動平均線的聚合程度',
        '技術綜合評分: 多項技術指標的加權平均'
    ]
    for feature in technical_features:
        p = doc.add_paragraph()
        p.add_run('- ' + feature)
    
    # 1.3 配置參數
    doc.add_heading('1.3 配置參數', 2)
    doc.add_paragraph('FeatureConfig 類別包含以下主要配置：')
    
    # 基礎配置
    doc.add_heading('基礎路徑設定：', 3)
    base_configs = [
        'BASE_DIR: "D:/Min/Python/Project/FA_Data"',
        'META_DATA_DIR: "meta_data"',
        'BACKUP_DIR: "backup"',
        'LOG_DIR: "logs"',
        'FEATURES_DIR: "features"'
    ]
    for config in base_configs:
        p = doc.add_paragraph()
        p.add_run('• ' + config)
    
    # 特徵參數
    doc.add_heading('特徵參數設定：', 3)
    feature_params = {
        '量能參數': [
            'short_period: 5',
            'long_period: 20',
            'volume_ma_periods: [5, 10, 20]',
            'volume_threshold: 2.0'
        ],
        '波動參數': [
            'short_period: 5',
            'long_period: 20',
            'std_window: 20',
            'atr_period: 14'
        ],
        '趨勢參數': [
            'ma_period: 20',
            'channel_period: 20',
            'momentum_periods: [5, 10, 20]',
            'trend_threshold: 0.02'
        ],
        '技術指標參數': [
            'rsi_period: 14',
            'ma_periods: [5, 10, 20, 60]',
            'macd_params: {fast: 12, slow: 26, signal: 9}',
            'kd_params: {k: 9, smooth_k: 3, smooth_d: 3}'
        ]
    }
    
    for category, params in feature_params.items():
        p = doc.add_paragraph()
        p.add_run(category + '：')
        for param in params:
            p = doc.add_paragraph()
            p.add_run('  - ' + param)
    
    # 保存文檔
    doc.save('D:/Min/Python/Project/note/tw_stock_analysis_guide.docx')

if __name__ == '__main__':
    create_doc()
