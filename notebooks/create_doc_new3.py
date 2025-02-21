# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def create_doc():
    doc = Document()
    
    # 添加標題
    title = doc.add_heading('特徵生成器(03_feature_generator)說明文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 功能概述
    doc.add_heading('1. 功能概述', 1)
    doc.add_paragraph('本特徵生成器用於處理台股數據，生成用於分析和預測的特徵。系統採用模塊化設計，支持多種特徵計算方法，並提供完整的數據驗證和錯誤處理機制。')
    
    # 1.1 核心功能
    doc.add_heading('1.1 核心功能', 2)
    core_features = [
        '量能特徵生成：計算與成交量相關的各類指標，包含量比、量增率等',
        '波動特徵計算：分析價格波動相關特徵，包含日內波動率、振幅等',
        '趨勢特徵分析：計算價格趨勢相關指標，包含均線系統、趨勢強度等',
        '技術指標整合：整合各類技術分析指標，如KD、RSI、MACD等',
        '特徵品質監控：監控特徵生成質量，包含數據完整性和有效性檢查',
        '自動報告生成：自動生成特徵分析報告，包含異常檢測和數據統計'
    ]
    for feature in core_features:
        doc.add_paragraph('• ' + feature)
    
    # 2. 系統架構
    doc.add_heading('2. 系統架構', 1)
    doc.add_paragraph('系統採用模塊化設計，各組件職責明確，便於維護和擴展。主要包含特徵管理、數據處理、特徵生成和驗證等模塊。')
    
    # 2.1 核心組件
    doc.add_heading('2.1 核心組件', 2)
    components = [
        'FeatureManager：特徵管理器，負責協調各個特徵生成器的運作\n' + 
        '  - 管理特徵生成流程\n' + 
        '  - 協調各個組件間的數據流轉\n' + 
        '  - 處理異常情況和錯誤恢復',
        
        'FeatureGenerator：基礎特徵生成器，提供通用的特徵計算方法\n' + 
        '  - 實現各類技術指標計算\n' + 
        '  - 提供數據預處理功能\n' + 
        '  - 支持自定義特徵擴展',
        
        'FeatureValidator：特徵驗證器，確保生成特徵的品質和完整性\n' + 
        '  - 數據完整性檢查\n' + 
        '  - 特徵有效性驗證\n' + 
        '  - 生成驗證報告',
        
        'DataLoader：數據加載器，負責讀取和預處理原始數據\n' + 
        '  - 支持多種數據源格式\n' + 
        '  - 處理缺失值和異常值\n' + 
        '  - 數據格式轉換和標準化'
    ]
    for component in components:
        doc.add_paragraph('• ' + component)
    
    # 2.2 數據流程
    doc.add_heading('2.2 數據流程', 2)
    doc.add_paragraph('系統數據處理流程如下：')
    flow_steps = [
        '1. 數據加載：從不同來源讀取原始數據',
        '2. 數據預處理：處理缺失值、異常值和格式轉換',
        '3. 特徵生成：計算各類技術指標和特徵',
        '4. 特徵驗證：驗證特徵的完整性和有效性',
        '5. 結果輸出：保存特徵數據和生成報告'
    ]
    for step in flow_steps:
        doc.add_paragraph('• ' + step)
    
    # 3. 特徵計算方法
    doc.add_heading('3. 特徵計算方法', 1)
    
    # 3.1 技術指標計算
    doc.add_heading('3.1 技術指標計算', 2)
    doc.add_paragraph('系統支持以下技術指標的計算：')
    
    tech_indicators = [
        'KD指標：\n' +
        '  - 計算方法：使用9日RSV作為基礎\n' +
        '  - 參數設置：K值和D值的平滑期數可配置\n' +
        '  - 應用場景：判斷超買超賣',
        
        'RSI指標：\n' +
        '  - 計算方法：基於價格變動的動能指標\n' +
        '  - 參數設置：支持多週期（如6、12、24日）\n' +
        '  - 應用場景：判斷價格動能',
        
        'MACD指標：\n' +
        '  - 計算方法：基於移動平均的趨勢指標\n' +
        '  - 參數設置：快線(12)、慢線(26)、DIF平滑(9)\n' +
        '  - 應用場景：判斷趨勢變化',
        
        '移動平均線：\n' +
        '  - 計算方法：支持SMA、EMA、WMA等\n' +
        '  - 參數設置：支持自定義週期\n' +
        '  - 應用場景：趨勢分析'
    ]
    for indicator in tech_indicators:
        doc.add_paragraph('• ' + indicator)
    
    # 3.2 特徵驗證方法
    doc.add_heading('3.2 特徵驗證方法', 2)
    validation_methods = [
        '數據完整性檢查：\n' +
        '  - 檢查缺失值比例\n' +
        '  - 驗證時間序列連續性\n' +
        '  - 檢查異常值',
        
        '特徵有效性驗證：\n' +
        '  - 檢查特徵值範圍\n' +
        '  - 驗證計算邏輯\n' +
        '  - 檢查特徵相關性',
        
        '數據質量報告：\n' +
        '  - 生成數據統計報告\n' +
        '  - 記錄異常情況\n' +
        '  - 提供改進建議'
    ]
    for method in validation_methods:
        doc.add_paragraph('• ' + method)
    
    # 4. 使用說明
    doc.add_heading('4. 使用說明', 1)
    
    # 4.1 環境配置
    doc.add_heading('4.1 環境配置', 2)
    env_setup = [
        '系統要求：\n' +
        '  - Python 3.8+\n' +
        '  - pandas 1.3+\n' +
        '  - numpy 1.20+\n' +
        '  - talib 0.4+',
        
        '依賴安裝：\n' +
        '  pip install -r requirements.txt',
        
        '配置文件：\n' +
        '  - config.py：主要配置文件\n' +
        '  - logging.conf：日誌配置\n' +
        '  - feature_params.json：特徵參數配置'
    ]
    for setup in env_setup:
        doc.add_paragraph('• ' + setup)
    
    # 4.2 使用示例
    doc.add_heading('4.2 使用示例', 2)
    usage = doc.add_paragraph()
    usage.add_run('基本使用示例：\n').bold = True
    usage.add_run('''
from feature_manager import FeatureManager

# 初始化特徵管理器
manager = FeatureManager()

# 設置時間範圍
start_date = '20230101'
end_date = '20231231'

# 生成特徵
manager.generate_features(start_date, end_date)

# 驗證特徵
manager.validate_features()

# 生成報告
manager.generate_report()
''')
    
    # 4.3 錯誤處理
    doc.add_heading('4.3 錯誤處理', 2)
    error_handling = [
        '常見錯誤：\n' +
        '  - 數據源不可用\n' +
        '  - 特徵計算異常\n' +
        '  - 驗證失敗',
        
        '處理方法：\n' +
        '  - 檢查數據源連接\n' +
        '  - 查看錯誤日誌\n' +
        '  - 調整參數配置',
        
        '錯誤恢復：\n' +
        '  - 自動重試機制\n' +
        '  - 數據備份恢復\n' +
        '  - 手動干預處理'
    ]
    for handling in error_handling:
        doc.add_paragraph('• ' + handling)
    
    # 5. 性能優化
    doc.add_heading('5. 性能優化', 1)
    performance_tips = [
        '數據處理優化：\n' +
        '  - 使用數據分片處理\n' +
        '  - 實現並行計算\n' +
        '  - 優化記憶體使用',
        
        '計算效率提升：\n' +
        '  - 使用向量化運算\n' +
        '  - 實現緩存機制\n' +
        '  - 優化算法邏輯',
        
        '系統配置優化：\n' +
        '  - 調整批處理大小\n' +
        '  - 優化日誌級別\n' +
        '  - 合理設置超時'
    ]
    for tip in performance_tips:
        doc.add_paragraph('• ' + tip)
    
    # 6. 注意事項
    doc.add_heading('6. 注意事項', 1)
    notes = [
        '數據源管理：\n' +
        '  - 確保數據源的完整性和準確性\n' +
        '  - 定期更新數據源\n' +
        '  - 備份重要數據',
        
        '特徵生成：\n' +
        '  - 注意特徵計算的時效性\n' +
        '  - 監控特徵質量\n' +
        '  - 及時處理異常',
        
        '系統維護：\n' +
        '  - 定期檢查系統日誌\n' +
        '  - 更新系統依賴\n' +
        '  - 優化系統配置'
    ]
    for note in notes:
        doc.add_paragraph('• ' + note)
    
    # 創建輸出目錄（如果不存在）
    output_dir = 'D:/Min/Python/Project/note'
    os.makedirs(output_dir, exist_ok=True)
    
    # 保存文檔
    doc.save(os.path.join(output_dir, 'tw_stock_analysis_guide.docx'))

if __name__ == '__main__':
    create_doc()
